from __future__ import annotations

import json
import logging
from pathlib import Path

from pypdf import PdfReader

from esg_rag.models import Document

logger = logging.getLogger(__name__)

_MIN_PAGE_CHARS = 60


class DocumentLoader:
    supported_extensions = {".txt", ".md", ".pdf", ".json", ".docx"}

    def load_directory(self, directory: Path) -> list[Document]:
        documents: list[Document] = []
        if not directory.exists():
            return documents

        for path in sorted(directory.rglob("*")):
            if path.is_file() and path.suffix.lower() in self.supported_extensions:
                docs = self.load_file(path)
                if docs:
                    logger.info("Loaded %d document(s) from %s", len(docs), path.name)
                else:
                    logger.warning("No content extracted from %s", path.name)
                documents.extend(docs)
        return documents

    def load_file(self, path: Path) -> list[Document]:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return [
                Document(
                    text=path.read_text(encoding="utf-8", errors="ignore"),
                    metadata=self._base_metadata(path),
                )
            ]
        if suffix == ".json":
            return self._load_json(path)
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix == ".docx":
            return self._load_docx(path)
        return []

    def _base_metadata(self, path: Path) -> dict[str, str]:
        return {
            "source": str(path),
            "source_name": path.name,
            "source_type": path.suffix.lower().lstrip("."),
        }

    def _load_json(self, path: Path) -> list[Document]:
        payload = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        if isinstance(payload, list):
            return [
                Document(
                    text=json.dumps(item, ensure_ascii=False, indent=2),
                    metadata={**self._base_metadata(path), "record_index": index},
                )
                for index, item in enumerate(payload)
            ]
        return [
            Document(
                text=json.dumps(payload, ensure_ascii=False, indent=2),
                metadata=self._base_metadata(path),
            )
        ]

    def _load_pdf(self, path: Path) -> list[Document]:
        if path.stat().st_size == 0:
            logger.warning("Skipping empty file: %s", path.name)
            return []
        try:
            reader = PdfReader(str(path))
        except Exception:
            logger.exception("Failed to open PDF: %s", path.name)
            return []

        base = self._base_metadata(path)
        total_pages = len(reader.pages)
        documents: list[Document] = []
        carry = ""
        carry_start: int | None = None

        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue

            if len(text) < _MIN_PAGE_CHARS:
                if carry_start is None:
                    carry_start = page_number
                carry += ("\n" if carry else "") + text
                continue

            if carry:
                text = carry + "\n" + text
                start_page = carry_start or page_number
                carry = ""
                carry_start = None
            else:
                start_page = page_number

            documents.append(Document(
                text=text,
                metadata={**base, "page": start_page, "total_pages": total_pages},
            ))

        if carry:
            if documents:
                last = documents[-1]
                last.text += "\n" + carry
            else:
                documents.append(Document(
                    text=carry,
                    metadata={**base, "page": carry_start or 1, "total_pages": total_pages},
                ))

        return documents

    def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument
        from docx.table import Table

        if path.stat().st_size == 0:
            logger.warning("Skipping empty file: %s", path.name)
            return []

        try:
            doc = DocxDocument(str(path))
        except Exception:
            logger.exception("Failed to open docx: %s", path.name)
            return []

        base = self._base_metadata(path)
        sections: list[tuple[str | None, list[str]]] = []
        current_heading: str | None = None
        current_paras: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "p":
                from docx.text.paragraph import Paragraph
                para = Paragraph(element, doc)
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading" in style_name or "title" in style_name:
                    if current_paras:
                        sections.append((current_heading, current_paras))
                    current_heading = text
                    current_paras = []
                else:
                    current_paras.append(text)

            elif tag == "tbl":
                table = Table(element, doc)
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    current_paras.append("\n".join(rows))

        if current_paras:
            sections.append((current_heading, current_paras))

        if not sections:
            return []

        documents: list[Document] = []
        chunk_size = 15

        for heading, paras in sections:
            for i in range(0, len(paras), chunk_size):
                batch = paras[i : i + chunk_size]
                text_parts = []
                if heading:
                    text_parts.append(f"## {heading}")
                text_parts.extend(batch)
                text = "\n\n".join(text_parts)
                if text.strip():
                    meta = {**base, "section_start": i}
                    if heading:
                        meta["section_heading"] = heading
                    documents.append(Document(text=text, metadata=meta))

        return documents
